# coding=utf-8
'''Power BI to docx report converter'''

import sys
import datetime
from openpyxl import load_workbook
from pathlib import Path
from subprocess import Popen

from docx.enum.text import  WD_ALIGN_PARAGRAPH, WD_LINE_SPACING \
                             # pylint: disable=E0611

from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, \
                            WD_ROW_HEIGHT_RULE # pylint: disable=E0611
                            
from docx import Document
from docx.shared import Cm, Pt

def load_sheet():

    xlsx_file = Path(sys.argv[1])
    wb_obj    = load_workbook(xlsx_file)
    sheet     = wb_obj.active
    return(sheet)    

def get_orginfo():

    sheet   = load_sheet()
    inn     = sheet['A4'].value
    kpp     = sheet['B4'].value
    orgname = sheet['C4'].value.split(sep=' ', maxsplit=1)[1]
    rawname = orgname.translate(str.maketrans('', '','.«»\'\"'))
    return([inn, kpp, orgname, rawname])
  
def prepare_template():

    orginfo = get_orginfo()
    inn     = orginfo[0]
    kpp     = orginfo[1]
    orgname = orginfo[2]
    start_date  = (datetime.date.today() - datetime.timedelta(days=365)). \
                                                   strftime('%d.%m.%Y')
    end_date    = datetime.date.today().strftime('%d.%m.%Y')

    with open('template.txt', 'r', encoding="utf-8") as template:
        template_text = template.readlines()

    paragraph_0 = template_text[0].replace('\n', '')
    paragraph_1 = template_text[1].replace('\n', '')
    paragraph_2 = template_text[2].replace('\n', '')

    
    paragraph_0 = paragraph_0.format(orgname, inn)
    paragraph_2 = paragraph_2.format(start_date, end_date)
    
    return([paragraph_0, paragraph_1, paragraph_2])
    
def prepare_document():

        document = Document('template.docx')

        following_text = prepare_template()

        document.sections[-1].top_margin    = Cm(1)
        document.sections[-1].bottom_margin = Cm(1)
        document.sections[-1].left_margin   = Cm(2)
        document.sections[-1].right_margin  = Cm(2)

        document.paragraphs[0].text = (following_text[0])
        document.add_paragraph(following_text[1])
        document.add_paragraph(following_text[2])

        for par in document.paragraphs:
            par.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
            par.paragraph_format.first_line_indent = Cm(1)
            par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            par.paragraph_format.space_before      = Pt(0)
            par.paragraph_format.space_after       = Pt(0)
            for run in par.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                
        return(document)

def prepare_data():

    sheet    = load_sheet()
    last_row = int(str(sheet.max_row))
    data_table = []
       
    for row in sheet.iter_rows(min_row=4, min_col=4, max_col=5,
                               max_row=last_row, values_only=True):
        data_table.append(list(row))

    for i in data_table:
        i[0] = datetime.datetime.strftime(i[0], '%d.%m.%Y, %H:%M:%S')

    return(data_table, len(data_table))
    
def table_format(table):
    '''Apply styles to various table elements'''

    table.style            = 'Table Grid'
    table.alignment        = WD_TABLE_ALIGNMENT.CENTER
    table.rows.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height   = Cm(0.8)

    for col in table.columns:

        for cell in col.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for par in cell.paragraphs:
                par.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.LEFT
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                par.paragraph_format.space_before      = Pt(0)
                par.paragraph_format.space_after       = Pt(0)

                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

    for cell in table.row_cells(0):

        for par in cell.paragraphs:
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in par.runs:
                run.font.bold = True

final_document = get_orginfo()[3] + '.docx'

document = prepare_document()

data = prepare_data()

result     = data[0]
table_size = data[1]

table = document.add_table(rows = table_size+1, cols = 2)
hdr_cells         = table.rows[0].cells
hdr_cells[0].text = 'Дата и время начала сеанса'
hdr_cells[1].text = 'Интернет-адрес рабочего места абонента'

table_cells = table._cells

for i in range(0, table_size):

    row_cells = table_cells[(i+1)*2:(i+2)*2]
    row_cells[0].text = result[i][0]
    row_cells[1].text = result[i][1]

table_format(table)

try:
    document.save(final_document)
except OSError as error_message:
    print("Error: {0} - {1}.".format(error_message.filename,
                                     error_message.strerror))

args = ['explorer', final_document]
Popen(args)